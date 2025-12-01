import os
import sys
import re
import tkinter as tk
from tkinter import messagebox, ttk
from pathlib import Path
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Pt,Cm, Inches
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

if getattr(sys, 'frozen', False):
    base_path = os.path.dirname(sys.executable)
else:
    base_path = os.path.dirname(__file__)

# Load Existing Buku Daftar excelbook
filename = os.path.join(base_path, "Buku_Daftar_UAT.xlsx")
logo_path = os.path.join(base_path, 'STRIDE Logo.png')
logo2_path = os.path.join(base_path, 'STRIDE Logo2.png')
workbook = load_workbook(filename)
sheet = workbook.active

# Determine the user's desktop path
desktop_path = str(Path.home() / "Desktop")

# Define paths for Word documents to be saved on the desktop
doc_out_path = os.path.join(desktop_path, 'doc_out_file')

# Ensure the directories exist
os.makedirs(doc_out_path, exist_ok=True)

# Find the existing table in the sheet
def get_table_range(sheet):
    for table in sheet.tables.values():
        return table.ref  # Returns the table range

# Find header of the column
column_indexes = {}
for col in sheet.iter_cols(1, sheet.max_column):
    header = col[0].value
    if header in ["REPORT NUMBER","TOT","NO. OF TEST","INTERNAL REFERENCE NUMBER","DATE RECEIVED","RECEIVED BY", "CONTACT PERSON","APPLICANT BY", "CLIENT","WORK TITLE", "QUANTITY", "SAMPLE MARKING", "WORK CLASS"]:
        column_indexes[header] = col[0].column  # Store the column index

# Open added_rows dictionary so we can start counting for our display table. This is important for proper doc output
added_rows = []

# Open workclass codes dictionary 'W01'
workclass_codes = {
    "MINDEF": "9230",
    "Berbayar": "9240",
    "Agensi Kerajaan": "9250",
    "STRIDE": "9260",
}

# Additional suffixes for Work Class MINDEF 'W02'
workclass_MINDEF_suffixes = ["D", "L", "U", "MAB"]

# Function to automatically generate reference number
def generate_reference_number(sheet, subgroup_selection, suffix=None):
    # Generate the INTERNAL REFERENCE NUMBER for the selected Work Class.
    # Get the corresponding code for the Work Class from the dictionary we open 'W01'
    code = workclass_codes.get(subgroup_selection)
    if not code:
        raise ValueError(f"Invalid Work Class: {subgroup_selection}")

    # Get the last two digits of the current year
    current_year_suffix = datetime.now().year % 100

    # Extract all reference numbers in the column
    reference_numbers = []
    if "INTERNAL REFERENCE NUMBER" in column_indexes:
        col_index = column_indexes["INTERNAL REFERENCE NUMBER"]
        for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=col_index, max_col=col_index):
            ref_number = row[0].value
            if ref_number and ref_number.startswith(f"PA/UAT/{code}/{current_year_suffix:02d}"):
                reference_numbers.append(ref_number)

    # Determine the next running number
    existing_numbers = [
        int(ref.split("/")[-1].split(" ")[0])  # Extract the running number
        for ref in reference_numbers
        if ref.split("/")[-1].split(" ")[0].isdigit()
    ]
    next_running_number = max(existing_numbers, default=0) + 1

    # Construct the reference number
    base_reference = f"PA/UAT/{code}/{current_year_suffix:02d}/{next_running_number:02d}"
    # For MINDEF, there's additonal suffix
    if suffix:
        base_reference += f" ({suffix})"

    return base_reference # return this base reference for future use

# We want a triggering event for workclass subgroup. I.e for when MINDEF and berbayar is selected
def on_workclass_subgroup(event):
    selected_class = workclass_var.get()
    if selected_class == "MINDEF": # if MINDEF is selected
        roman_dropdown["values"] = ["Bekalan", "Pembangunan Spesifikasi", "Penyiasatan", "Lain-lain"]
        roman_dropdown.config(state="readonly")  # ...Enable the dropdown
        roman_var.set("")  # Clear previous selection

        additional_dropdown["values"] = ["D", "U", "L", "MAB", "Others"]
        additional_dropdown.config(state="readonly")  # Enable additional dropdown for suffix
        additional_var.set("")  # Clear previous additional selection

    elif selected_class == "Berbayar": # if berbayar is selected
        roman_dropdown["values"] = ["Tender", "Syarikat"] # The value for the dropdown is different from MINDEF
        roman_dropdown.config(state="readonly")  # ...Enable the dropdown
        roman_var.set("")  # Clear previous selection

        additional_dropdown.config(state="disabled")  # Disable additional dropdown for Berbayar group
        additional_var.set("")  # Clear additional selection

    else:
        roman_dropdown.config(state="disabled")  # If Berbayar and MINDEF is not selected...disable
        roman_var.set("")  # Clear the value

        additional_dropdown.config(state="disabled")  # Disable additional dropdown for Berbayar
        additional_var.set("")  # Clear additional selection

# Function to process input from user
def process_input():
    report_number = entry_RN.get()
    contact_person = entry_CP.get()
    received_by = entry_Rby.get()
    date_received = entry_date_received.get()
    worktitle = entry_worktitle.get()
    applicant_by = entry_applicantby.get()
    client = entry_client.get()
    kuantiti = entry_kuantiti.get()
    marking_input = entry_sample_marking.get()
    workclass = workclass_var.get()
    subgroup_selection = roman_var.get()
    additional_selection = additional_var.get()
    workclass_MINDEF_suffix = additional_selection if workclass == "MINDEF" and additional_selection in workclass_MINDEF_suffixes else None

    # Test Mapping for our page 2 and page 3 docs
    test_mapping = {
        "I": microscopic_var,
        "II": burning_var,
        "III": solubility_var,
        "IV": quantitative_var,
        "V": ftir_var,
    }

    # Raise error if none of the field is filled. Do note that applicant is not listed. We create dict first...
    fields = {
        "Report Number": report_number,
        "Contact Person": contact_person,
        "Received by": received_by,
        "Work Title": worktitle,
        "Date Received": date_received,
        "Client": client,
        "Kuantiti": kuantiti,
        "Sample Marking": marking_input,
        "Work Class": workclass,
    }

    # ...so that it is less cumbersome and repetitive
    for field_name, field_value in fields.items():
        if not field_value.strip():
            raise ValueError(f"{field_name} cannot be empty.")

    # Validate Additional Selection Mindef and Berbayar
    if workclass in ["MINDEF", "Berbayar"] and not subgroup_selection:
        raise ValueError("You must select additional information for Work Class MINDEF or Berbayar.")
    if workclass == "MINDEF" and not additional_selection:
        raise ValueError("You must select an additional option (D,U,L,MAB or Others) for Work Class MINDEF.")

    # Format Work Class with suffix for Mindef and Berbayar if ada. And we want to save this in excel
    display_workclass = f"{workclass} ({subgroup_selection})" if subgroup_selection else workclass

    # Generate the INTERNAL REFERENCE NUMBER by calling back function @ Line 65
    internal_reference_number = generate_reference_number(sheet, workclass, workclass_MINDEF_suffix)

    # Process Marking input by looking for commas. We call back input @ line 137 and process it to new name.
    # The name of processed marking input that we processes is processed items...😂
    items = marking_input.split(",")
    processed_items = []

    # Counter for numbering items
    index = 1

    for item in items:
        # strip whitespace from marking_input entry [cth: " Shirt" space before S is whitespace
        item = item.strip()
        # if there's nothing left after stripping (end of list)
        if not item:
            continue

            # Split into name and quantity based on ";"
        if ";" in item:
            name, qty = item.split(";", 1)  # Split into two parts only
            name = name.strip()  # Clean up whitespace around the name
            qty = qty.strip()  # Clean up whitespace around the quantity
            qty = int(qty) if qty.isdigit() else 1  # Default quantity to 1 if not provided or invalid
        else:
            name = item
            qty = 1  # Default quantity to 1 if ";" is missing

        # Format the index to always be two digits. This helps with our page 1 populate function
        formatted_index = str(index).zfill(2)

        processed_items.append(f"{formatted_index}. {name.lower()} x {qty}")

        # Increment the counter for the next item
        index += 1

    # Combine processed items into a single string for saving
    processed_marking = "; ".join(processed_items)

    # Process selected checkboxes. To count the number of test and what test is selected (T.O.T and N.O.T)
    selected_tests = [code for code, var in test_mapping.items() if var.get() == 1]
    number_of_test = len(selected_tests)

    # Join selected tests as a comma-separated string to save in TOT
    type_of_testing = ", ".join(selected_tests) if selected_tests else "None"

    # This return order has to be same in the order we re-called later in save_data function
    return date_received, client, kuantiti, report_number, received_by, applicant_by, worktitle, subgroup_selection, contact_person, processed_marking, display_workclass, additional_selection, internal_reference_number, type_of_testing, number_of_test

# Function to save data to Excel
def save_data():
    try:
        # Process inputs and get formatted data. Get from process_input function. Line 226
        date_received, client, kuantiti, report_number, received_by, applicant_by, worktitle, subgroup_selection, contact_person, processed_marking, display_workclass, additional_selection, internal_reference_number, type_of_testing, number_of_test = process_input()

        # Get table range from defined function
        table_range = get_table_range(sheet)

        # Find the next available row
        next_row = sheet.max_row + 1

        # Get "Applicant By" value or set to "NA" if empty
        applicant_by = entry_applicantby.get().strip() or "NA"

        # Insert the data into the corresponding columns in excel
        if "DATE RECEIVED" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["DATE RECEIVED"], value=date_received)
        if "CLIENT" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["CLIENT"], value=client)
        if "TOT" in column_indexes:  # Save the type of testing in the TOT column
            sheet.cell(row=next_row, column=column_indexes["TOT"], value=type_of_testing)
        if "NO. OF TEST" in column_indexes:  # Save the type of testing in the TOT column
            sheet.cell(row=next_row, column=column_indexes["NO. OF TEST"], value=number_of_test)
        if "WORK TITLE" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["WORK TITLE"], value=worktitle)
        if "QUANTITY" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["QUANTITY"], value=kuantiti)
        if "SAMPLE MARKING" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["SAMPLE MARKING"], value=processed_marking)
        if "REPORT NUMBER" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["REPORT NUMBER"], value=report_number)
        if "CONTACT PERSON" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["CONTACT PERSON"], value=contact_person)
        if "RECEIVED BY" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["RECEIVED BY"], value=received_by)
        if "APPLICANT BY" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["APPLICANT BY"], value=applicant_by)
        if "WORK CLASS" in column_indexes:  # Add Work Class to the appropriate column
            sheet.cell(row=next_row, column=column_indexes["WORK CLASS"], value=display_workclass)
        if "INTERNAL REFERENCE NUMBER" in column_indexes:
            sheet.cell(row=next_row, column=column_indexes["INTERNAL REFERENCE NUMBER"], value=internal_reference_number)

        # Apply Excel-specific formatting for DATE, Quantity and No. of Test
        if "DATE RECEIVED" in column_indexes:
            date_col = column_indexes["DATE RECEIVED"]
            for row in range(2, next_row + 1):  # Assuming headers are in row 1
                sheet.cell(row=row, column=date_col).number_format = "dd/mm/yy"

        if "QUANTITY" in column_indexes:
            quantity_col = column_indexes["QUANTITY"]
            for row in range(2, next_row + 1):  # Assuming headers are in row 1
                sheet.cell(row=row, column=quantity_col).number_format = "General"

        if "NO. OF TEST" in column_indexes:
            test_col = column_indexes["NO. OF TEST"]
            for row in range(2, next_row + 1):  # Assuming headers are in row 1
                sheet.cell(row=row, column=test_col).number_format = "General"

        # Update new table range by appending new row number
        end_row = next_row
        new_range = f"A1:{sheet.cell(row=end_row, column=sheet.max_column).coordinate}"

        # Entered the updated table range into the existing table range
        if table_range:
            table = sheet.tables[sheet.tables.keys().__iter__().__next__()]  # Only one table here
            table.ref = new_range  # Update the table reference

        # Save the updated workbook
        workbook.save(filename)

        # Update status label to show success message
        label_status.config(text="Data saved successfully!")

        # keep track of each new additonal row number
        added_rows.append(next_row)

        # Ask if the user wants to add another entry
        result = messagebox.askyesno("Add More Data", "Do you want to add another entry?")

        if result:  # If "Yes" is clicked. We clear all entry on tkinter
            entry_date_received.delete(0, tk.END)  # Clear the entry fields
            entry_client.delete(0, tk.END)
            entry_kuantiti.delete(0, tk.END)
            entry_sample_marking.delete(0, tk.END)
            entry_RN.delete(0, tk.END)
            entry_CP.delete(0, tk.END)
            entry_Rby.delete(0, tk.END)
            entry_worktitle.delete(0, tk.END)
            for var in [microscopic_var,burning_var,solubility_var,quantitative_var,ftir_var]:
                var.set(0)
            workclass_var.set("")
            roman_var.set("")
            additional_var.set("")
        else:  # If "No" is clicked
            # Destroy the current window and open the new table window
            root.quit()  # Exit the main event loop
            root.destroy()  # Destroy the Tkinter window

            # Call display_table() to show the data in a new window
            display_table(next_row)  # Call the function to display the table

    except ValueError as e:
        messagebox.showerror("Error", str(e))

# Start with common doc function font setting, spacing, table, merge etc2
def font_settings_header(run, font_name='Arial', font_size=Pt(11), bold=True, underline=False):
    run.font.name = font_name
    run.font.size = font_size
    run.bold = bold
    if underline:
        run.font.underline = WD_UNDERLINE.SINGLE  # Use WD_UNDERLINE for underlining
    else:
        run.font.underline = None  # No underline

def set_paragraph_spacing(paragraph, line_spacing_pt, before_spacing_pt=0, after_spacing_pt=0):
    pPr = paragraph._element.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), str(int(line_spacing_pt * 20)))
    spacing.set(qn('w:lineRule'), 'auto')

    # Set before and after spacing
    spacing.set(qn('w:before'), str(int(before_spacing_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_spacing_pt * 20)))

# Function to set cell text
def set_cell_text(cell, text, bold=False, font_size = 10, line_spacing_pt=11, before_spacing_pt=0, after_spacing_pt=0, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.font.name = 'Arial'  # Set font name
    run.font.size = Pt(font_size)  # Set font size
    if bold:
        run.bold = True
    paragraph.alignment = alignment
    set_paragraph_spacing(paragraph, line_spacing_pt, before_spacing_pt, after_spacing_pt)

def set_column_width(column, width):
    for cell in column.cells:
        tc = cell._element  # Access the XML element of the cell
        tcPr = tc.get_or_add_tcPr()  # Get or create the cell properties element
        tcW = OxmlElement('w:tcW')  # Create a new width element
        tcW.set(qn('w:w'), str(int(width * 1440)))  # 1440 twips per inch, set the width
        tcW.set(qn('w:type'), 'dxa')  # Use 'dxa' for width units in twips (20ths of a point)
        tcPr.append(tcW)

def set_table_borders(table):
    tbl = table._tbl  # Get the table XML element
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))

    if tblBorders is None:
       tblBorders = OxmlElement('w:tblBorders')
       tblPr.append(tblBorders)

        # Define border styles
    border_styles = {
        'top': {'w:val': 'single', 'w:sz': '4'},
        'left': {'w:val': 'single', 'w:sz': '4'},
        'bottom': {'w:val': 'single', 'w:sz': '4'},
        'right': {'w:val': 'single', 'w:sz': '4'},
        'insideH': {'w:val': 'single', 'w:sz': '4'},
        'insideV': {'w:val': 'single', 'w:sz': '4'}
    }

    # Apply border styles to each side of the table and inside
    for side, attrs in border_styles.items():
        border_element = OxmlElement(f'w:{side}')
        for key, value in attrs.items():
            border_element.set(qn(key), value)
        tblBorders.append(border_element)


def merge_cells_horizontally(table, row_idx, start_col_idx, end_col_idx):
    cell = table.cell(row_idx, start_col_idx)
    cell.merge(table.cell(row_idx, end_col_idx))


# Function to create a Word document for a specific row
def create_page1(row_num):
    # Call all data we needed first
    report_number = sheet.cell(row=row_num, column=column_indexes["REPORT NUMBER"]).value
    client = sheet.cell(row=row_num, column=column_indexes["CLIENT"]).value
    contact_person = sheet.cell(row=row_num, column=column_indexes["CONTACT PERSON"]).value
    date_received = sheet.cell(row=row_num, column=column_indexes["DATE RECEIVED"]).value
    received_by = sheet.cell(row=row_num, column=column_indexes["RECEIVED BY"]).value
    work_title = sheet.cell(row=row_num, column=column_indexes["WORK TITLE"]).value
    sample_marking = sheet.cell(row=row_num, column=column_indexes["SAMPLE MARKING"]).value
    workclass = sheet.cell(row=row_num, column=column_indexes["WORK CLASS"]).value
    lab_work_no = sheet.cell(row=row_num, column=column_indexes["INTERNAL REFERENCE NUMBER"]).value
    applicant_by = sheet.cell(row=row_num, column=column_indexes["APPLICANT BY"]).value
    type_of_test = sheet.cell(row=row_num, column=column_indexes["TOT"]).value

    # Split the string into individual items. For sample marking
    items = sample_marking.split("; ")

    # Extract names from each item. We sort it so there's no trouble at all...NO TROUBLE AT ALL...
    # Okay...there's a minor bug...but other than that, there's no trouble at all...
    #...sigh!...We need to figure out a sort that would really follow the order user type it in.
    # here we sort the item so it would populate in the order of entry by the user
    marking_for_table = sorted(item[3:-4].strip().title() for item in items)

    # Additional function we create just specific for page one
    def apply_single_line_spacing_to_table(table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_spacing(paragraph, line_spacing_pt=11, before_spacing_pt=1, after_spacing_pt=1)

    def set_cell_border(cell, is_bold=True, Thickness=6, borders='top and bottom'):
        # Get or create the cell properties element (w:tcPr)
        tc_pr = cell._element.get_or_add_tcPr()

        # Get or create the borders element (w:tcBorders)
        cell_borders = tc_pr.find(qn('w:tcBorders'))
        if cell_borders is None:
            cell_borders = OxmlElement('w:tcBorders')
            tc_pr.append(cell_borders)

        # Set top border if 'top_and_bottom' or 'top_only'
        if borders in ['top_and_bottom', 'top_only']:
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')  # Solid line
            top_border.set(qn('w:sz'), str(Thickness))  # Thickness (in 1/8 pt)
            top_border.set(qn('w:space'), '0')  # No extra space
            if is_bold:
                top_border.set(qn('w:color'), '000000')  # Black color for bold effect
            cell_borders.append(top_border)

        # Set bottom border if 'top_and_bottom' or 'bottom_only'
        if borders in ['top_and_bottom', 'bottom_only']:
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')  # Solid line
            bottom_border.set(qn('w:sz'), str(Thickness))  # Thickness (in 1/8 pt)
            bottom_border.set(qn('w:space'), '0')  # No extra space
            if is_bold:
                bottom_border.set(qn('w:color'), '000000')  # Black color for bold effect
            cell_borders.append(bottom_border)

        # Set left and right borders to 'nil' (no border)
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'nil')  # No left border
        cell_borders.append(left_border)

        right_border = OxmlElement('w:right')
        right_border.set(qn('w:val'), 'nil')  # No right border
        cell_borders.append(right_border)

    # This is the populate table function for line 430...the one with "NO TROUBLE AT ALL"
    def populate_table(name, table, max_items_per_column=12, start_row=1):
        total_columns = 4  # Total number of columns in the table
        flat_columns = list(range(total_columns))  # Flatten the columns for straightforward indexing

        # Populate the table
        for idx, item in enumerate(name):
            # Determine the column and row for this item
            current_column = (idx // max_items_per_column) % total_columns  # Cycle through columns sequentially
            col_idx = flat_columns[current_column]  # Get the column index
            row_idx = start_row + (idx // (total_columns * max_items_per_column)) * max_items_per_column + (
                    idx % max_items_per_column)

            # Ensure row index is valid within table boundaries
            if row_idx >= len(table.rows):
                raise ValueError(f"Table has insufficient rows to fit all items starting from row {start_row}.")

            # Set the cell text
            cell = table.cell(row_idx, col_idx)
            cell.text = f"{idx + 1}. {item}"

            # Set the font to Arial, size 10
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)

        return table

    # Create a new document
    doc = Document()

    # Set the top, bottom, left, and right margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)  # Set the top margin to 0 cm
        section.bottom_margin = Cm(2.25)  # Set the bottom margin to 2.25 cm
        section.left_margin = Cm(2.03)  # Set the left margin to 2.03 cm
        section.right_margin = Cm(2.03)  # Set the right margin to 2.03 cm

    # Add the first line: "Nama Unit"
    p1 = doc.add_paragraph()
    p1.alignment = 1  # Center alignment
    p1.paragraph_format.space_before = Pt(12)  # 12 pt is roughly one line of spacing
    run1 = p1.add_run("UNIT ANALISIS TEKSTIL")
    font_settings_header(run1, bold=True, underline=False)
    set_paragraph_spacing(p1, line_spacing_pt=11, before_spacing_pt=0, after_spacing_pt=0)

    # Add the second line: "Nama Bahagian"
    p2 = doc.add_paragraph()
    p2.alignment = 1  # Center alignment
    run2 = p2.add_run("BAHAGIAN TEKNOLOGI PRESTASI ANGKATAN")
    font_settings_header(run2, bold=True, underline=False)
    set_paragraph_spacing(p2, line_spacing_pt=11, before_spacing_pt=0, after_spacing_pt=0)

    # Add Space antara Nama bahagian dan document title
    p4 = doc.add_paragraph()
    p4.alignment = 1  # Center alignment
    run4 = p4.add_run()
    set_paragraph_spacing(p4,line_spacing_pt=11, before_spacing_pt=0, after_spacing_pt=0)

    # Add the third line: "Document Title"
    p3 = doc.add_paragraph()
    p3.alignment = 1  # Center alignment
    run3 = p3.add_run("Laboratory Work File")
    font_settings_header(run3, bold=True, underline=False)

    # Create table with specific number of rows and columns
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Disable auto-fit so the widths are respected
    table.autofit = False

    set_table_borders(table)

    apply_single_line_spacing_to_table(table)

    set_column_width(table.columns[0], 1.12)
    set_column_width(table.columns[1], 3.26)
    set_column_width(table.columns[2], 1.15)
    set_column_width(table.columns[3], 1.65)

    merge_cells_horizontally(table, 5, 1, 3)  # Merge Row Work Title
    merge_cells_horizontally(table, 4, 1, 3)  # Merge Row Specification

    set_cell_text(table.cell(0, 0), 'Report No.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 0), 'Work Class', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(2, 0), 'Client', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(3, 0), 'Contact Person', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(4, 0), 'Work Title', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(5, 0), 'Specification', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(0, 2), 'Lab Work No.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 2), 'Applicant by', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(2, 2), 'Date Received', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(3, 2), 'Received by', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Content populate from Excel
    set_cell_text(table.cell(0, 1), report_number, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 1), workclass, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(2, 1), client, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(3, 1), contact_person, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(4, 1), work_title, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(0, 3), lab_work_no, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 3), applicant_by, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(2, 3), date_received, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(3, 3), received_by, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Add the fourth line: "Sample Marking"
    p4 = doc.add_paragraph()
    p4.alignment = 1  # Center alignment
    run4 = p4.add_run("Sample Marking")
    font_settings_header(run4,font_size=Pt(10),bold=False, underline=False)
    set_paragraph_spacing(p4, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3)

    # Create table with specific number of rows and columns
    table01 = doc.add_table(rows=13, cols=4)
    table01.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Disable auto-fit so the widths are respected
    table01.autofit = False

    set_table_borders(table01)

    set_column_width(table01.columns[0], 1.8)
    set_column_width(table01.columns[1], 1.8)
    set_column_width(table01.columns[2], 1.8)
    set_column_width(table01.columns[3], 1.8)

    populate_table(marking_for_table, table01, max_items_per_column=12, start_row=1)

    apply_single_line_spacing_to_table(table01)

    # Add the fifth line:
    p5 = doc.add_paragraph()
    p5.alignment = 0  # Center alignment
    run5 = p5.add_run("Laboratory Activities:")
    font_settings_header(run5, font_size=Pt(10), bold=True, underline=False)
    set_paragraph_spacing(p5, line_spacing_pt=11, before_spacing_pt=12, after_spacing_pt=3)

    # Create table with specific number of rows and columns (number of rows at top, olah dari json)
    table02 = doc.add_table(rows=4, cols=5)
    table02.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Disable auto-fit so the widths are respected
    table02.autofit = False

    set_table_borders(table02)

    set_cell_text(table02.cell(0, 0), 'No.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(0, 1), 'Type of Testing', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(0, 2), 'Laboratory Personnel', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(0, 3), 'Date Start', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(0, 4), 'Date Finish', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Start populating the table based on type_of_test
    row_index = 1  # Starting row index for data

    # Check if any value from {I, II, III, IV} is present in type_of_test
    if any(test in type_of_test for test in {'I', 'II', 'III', 'IV'}):
        table02.add_row()
        set_cell_text(table02.cell(row_index, 0), '1.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table02.cell(row_index, 1), 'Qualitative Analysis', line_spacing_pt=11, before_spacing_pt=3,
                      after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        row_index += 1

    # Check if {IV} is present in type_of_test
    if 'IV' in type_of_test:
        table02.add_row()
        set_cell_text(table02.cell(row_index, 0), '2.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table02.cell(row_index, 1), 'Quantitative Analysis', line_spacing_pt=11, before_spacing_pt=3,
                      after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)

    set_column_width(table02.columns[0], 0.37)
    set_column_width(table02.columns[1], 2.32)
    set_column_width(table02.columns[2], 1.66)
    set_column_width(table02.columns[3], 1.42)
    set_column_width(table02.columns[4], 1.42)

    apply_single_line_spacing_to_table(table02)

    # Add the sixth line
    p6 = doc.add_paragraph()
    p6.alignment = 0  # Center alignment
    run6 = p6.add_run("Report:")
    font_settings_header(run6, font_size=Pt(10), bold=True, underline=False)
    set_paragraph_spacing(p6, line_spacing_pt=11, before_spacing_pt=12, after_spacing_pt=3)

    # Create table with specific number of rows and columns (number of rows at top, olah dari json)
    table03 = doc.add_table(rows=5, cols=2)
    table03.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Disable auto-fit so the widths are respected
    table03.autofit = False

    set_table_borders(table03)

    apply_single_line_spacing_to_table(table03)

    set_column_width(table03.columns[0], 1.91)
    set_column_width(table03.columns[1], 5.28)

    set_cell_text(table03.cell(0, 0), 'Report No.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(1, 0), 'Report Title', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(2, 0), 'Draft Date', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(3, 0), 'Approved Date', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(4, 0), 'Dispatch Date', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(0, 1), report_number, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(1, 1), work_title, line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Add the seventh line
    p7 = doc.add_paragraph()
    p7.alignment = 0  # Center alignment
    run7 = p7.add_run("Remarks:")
    font_settings_header(run7, font_size=Pt(10), bold=False, underline=False)
    set_paragraph_spacing(p7, line_spacing_pt=11, before_spacing_pt=12, after_spacing_pt=3)

    # Adding double parallel bold lines
    table04 = doc.add_table(rows=3, cols=1)

    apply_single_line_spacing_to_table(table04)

    # Access the cell in the table (but don't add any text)
    Is_bold = False
    border_thickness = 3
    for row in table04.rows:
        for cell in row.cells:
            set_cell_border(cell, is_bold=Is_bold, Thickness=border_thickness, borders='bottom_only')

    # Save the document
    work_file = os.path.join(doc_out_path,f'{client}_workfile.docx')
    doc.save(work_file)
    messagebox.showinfo("Document Created", f"Word document '{work_file}' created successfully.")

def create_page2(row_num):
    # Call data locally for easy access
    work_title = sheet.cell(row=row_num, column=column_indexes["WORK TITLE"]).value
    sample_marking = sheet.cell(row=row_num, column=column_indexes["SAMPLE MARKING"]).value
    workclass = sheet.cell(row=row_num, column=column_indexes["WORK CLASS"]).value
    lab_work_no = sheet.cell(row=row_num, column=column_indexes["INTERNAL REFERENCE NUMBER"]).value
    type_of_test = sheet.cell(row=row_num, column=column_indexes["TOT"]).value
    client = sheet.cell(row=row_num, column=column_indexes["CLIENT"]).value

    # Map type of test so that we know which cell to mark check or cross appropriately
    test_to_cell_mapping = {
        'I': (1, 1),
        'II': (0, 1),
        'III': (2, 1),
        'IV': (3, 1),
        'V': (4, 1)
    }

    # Split the string into individual items. Same old splitting and regurgitate
    items = sample_marking.split("; ")

    # Extract names from each item. Again we sort. And this time. I think it is unironically no trouble at all
    marking_for_table = sorted(item[2:-4].strip().title() for item in items)

    def populate_names_in_table(cell, names):
        num_names = len(names)

        # Determine number of columns based on number of names. Max we have 3 column, max item 36
        if 10 < num_names <= 20:
            num_columns = 2
            names_per_column = 10
        elif 20 < num_names <= 30:
            num_columns = 3
            names_per_column = 10
        else:
            num_columns = 3
            names_per_column = 12

        # Split names into chunks
        columns = [names[i:i + names_per_column] for i in range(0, num_names, names_per_column)]

        # Add a single table within the cell
        table = cell.add_table(rows=1, cols=num_columns)
        table.autofit = False  # Disable autofit to prevent unnecessary column resizing

        index = 1
        for col_idx, column_names in enumerate(columns):
            # Fill each column with names
            col_cell = table.cell(0, col_idx)
            for name in column_names:
                para = col_cell.add_paragraph()
                run = para.add_run(f"{index}. {name}")  # Add index and name
                run.font.name = "Arial"  # Set font to Arial
                run.font.size = Pt(9)  # Set font size to 9
                para.paragraph_format.line_spacing = Pt(10)  # Minimize spacing between lines
                para.paragraph_format.space_after = Pt(0)  # Remove extra spacing
                index += 1

        # Remove internal borders. So there's no table between the column for marking names
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                tc_borders = OxmlElement('w:tcBorders')
                for border in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border}')
                    border_elem.set(qn('w:val'), 'nil')  # Correctly set the 'w:val' attribute
                    tc_borders.append(border_elem)
                tc_pr.append(tc_borders)

    # Create a new document
    doc = Document()

    # Set the top, bottom, left, and right margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)  # Set the top margin to 0 cm
        section.bottom_margin = Inches(0.39)  # Set the bottom margin to 2.25 cm
        section.left_margin = Inches(1.0)  # Set the left margin to 2.03 cm
        section.right_margin = Inches(1.0)  # Set the right margin to 2.03 cm

    header = doc.sections[0].header  # Get the first section of the document
    header_table = header.add_table(rows=2, cols=3,width=Inches(6)) # Access the header of the section
    header_table.autofit = False

    # Set explicit column widths
    column_widths = [Inches(1.9), Inches(3.1), Inches(1.5)]  # Adjust as needed
    for col, width in zip(header_table.columns, column_widths):
        for cell in col.cells:
            cell.width = width

    set_table_borders(header_table)

    # Add the logo to the first cell
    cell = header_table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(1.1))

    # Merge the first column (2 rows)
    cell_to_merge = header_table.cell(1, 0)
    cell.merge(cell_to_merge)

    # Add content to the other cells
    set_cell_text(header_table.cell(0, 1), 'WORK SHEET', bold=True, line_spacing_pt=12, before_spacing_pt=6, after_spacing_pt=6)
    set_cell_text(header_table.cell(1, 1), '\n\nLaboratory Testing Form', bold=True, line_spacing_pt=12, before_spacing_pt=0, after_spacing_pt=0)
    set_cell_text(header_table.cell(0, 2), 'Document No: STRIDE/TAL/WS/01', font_size=9, bold=False, line_spacing_pt=12, before_spacing_pt=2, after_spacing_pt=2,alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(header_table.cell(1, 2), '\nPage: 1 of 1\nIssue No: 1\nRev. No: 0\nIssue Date: 1/3/2021', font_size=9, bold=False, line_spacing_pt=12, before_spacing_pt=2, after_spacing_pt=2,alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_column_width(table.columns[0],2.44)
    set_column_width(table.columns[1], 3.84)

    set_cell_text(table.cell(0, 0), 'Internal Reference', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6,alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 0), 'Work Title', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(2, 0), 'Sample Marking', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(3, 0), 'Specification', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(4, 0), 'Test Result', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(5, 0), 'Work Class', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    set_cell_text(table.cell(0, 1), lab_work_no, bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 1), work_title, bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(5, 1), workclass, bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    populate_names_in_table(table.cell(2,1),marking_for_table)

    p1 = doc.add_paragraph()
    run1 = p1.add_run("   General Information")
    font_settings_header(run1, bold=True, underline=False)
    p1.paragraph_format.space_before = Pt(3)
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p2 = doc.add_paragraph()
    run2 = p2.add_run("   Testing")
    font_settings_header(run2, bold=True, underline=False)
    p2.paragraph_format.space_after = Pt(3)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table01 = doc.add_table(rows=5, cols=2)
    table01.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table01)
    set_column_width(table01.columns[0], 2.44)
    set_column_width(table01.columns[1], 3.86)

    set_table_borders(table01)

    set_cell_text(table01.cell(0, 0), 'Qualitative Analysis (Burning)', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table01.cell(1, 0), 'Qualitative Analysis (Microscopic)', bold=False, line_spacing_pt=12,
                  before_spacing_pt=6,after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table01.cell(2, 0), 'Qualitative Analysis (Solubility)', bold=False, line_spacing_pt=12,
                  before_spacing_pt=6,after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table01.cell(3, 0), 'Quantitative Analysis', bold=False, line_spacing_pt=12,
                  before_spacing_pt=6,after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table01.cell(4, 0), 'FTIR', bold=False, line_spacing_pt=12,
                  before_spacing_pt=6,after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    for test, cell_coords in test_to_cell_mapping.items():
        row, col = cell_coords
        if test in type_of_test:  # Assuming 'type_of_test' is defined and contains the test types
            set_cell_text(table01.cell(row, col),'√',bold=False,line_spacing_pt=12,before_spacing_pt=6,
                          after_spacing_pt=6,alignment=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            set_cell_text(table01.cell(row, col),'X',bold=False,line_spacing_pt=12,before_spacing_pt=6,
                after_spacing_pt=6,alignment=WD_ALIGN_PARAGRAPH.CENTER)

    p3 = doc.add_paragraph()
    run3 = p3.add_run("   Carried out by:")
    font_settings_header(run3, bold=True, underline=False)
    p3.paragraph_format.space_before = Pt(6)
    p3.paragraph_format.space_after = Pt(3)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table02 = doc.add_table(rows=3, cols=2)
    table02.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_column_width(table02.columns[0], 1.41)
    set_column_width(table02.columns[1], 4.82)

    set_cell_text(table02.cell(0, 1)," : "+'.'*50, bold=False, line_spacing_pt=12,
                  before_spacing_pt=6,after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(1, 0), 'Date Start', bold=False, line_spacing_pt=12,
                  before_spacing_pt=0,after_spacing_pt=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(2, 0), 'Date Completed', bold=False, line_spacing_pt=12,
                  before_spacing_pt=0, after_spacing_pt=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(1, 1), " : " + '.' * 50, bold=False, line_spacing_pt=12,
                  before_spacing_pt=6, after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(2, 1), " : " + '.' * 50, bold=False, line_spacing_pt=12,
                  before_spacing_pt=6, after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    p4 = doc.add_paragraph()
    run4 = p4.add_run("   Verified by:")
    font_settings_header(run4, bold=True, underline=False)
    p4.paragraph_format.space_before = Pt(6)
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table03 = doc.add_table(rows=2, cols=2)
    table03.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_column_width(table03.columns[0], 1.41)
    set_column_width(table03.columns[1], 4.82)

    set_cell_text(table03.cell(0, 1), " : " + '.' * 50, bold=False, line_spacing_pt=12,
                  before_spacing_pt=6, after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(1, 0), 'Date', bold=False, line_spacing_pt=12,
                  before_spacing_pt=0, after_spacing_pt=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table03.cell(1, 1), " : " + '.' * 50, bold=False, line_spacing_pt=12,
                  before_spacing_pt=6, after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Save the document
    test_form = os.path.join(doc_out_path, f'{client}_testform.docx')
    doc.save(test_form)
    messagebox.showinfo("Document Created", f"Word document '{test_form}' created successfully.")

def create_page3(row_num):
    # Call locally data for easy access
    work_title = sheet.cell(row=row_num, column=column_indexes["WORK TITLE"]).value
    lab_work_no = sheet.cell(row=row_num, column=column_indexes["INTERNAL REFERENCE NUMBER"]).value
    type_of_test = sheet.cell(row=row_num, column=column_indexes["TOT"]).value
    client = sheet.cell(row=row_num, column=column_indexes["CLIENT"]).value

    # Create a new document
    doc = Document()

    # Set the top, bottom, left, and right margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)  # Set the top margin to 0 cm
        section.bottom_margin = Inches(0.39)  # Set the bottom margin to 2.25 cm
        section.left_margin = Inches(1.0)  # Set the left margin to 2.03 cm
        section.right_margin = Inches(1.0)  # Set the right margin to 2.03 cm

    header = doc.sections[0].header  # Get the first section of the document
    header_table = header.add_table(rows=1, cols=2, width=Inches(6))  # Access the header of the section
    header_table.autofit = False

    # Set explicit column widths
    column_widths = [Inches(1.1), Inches(5.7)]  # Adjust as needed
    for col, width in zip(header_table.columns, column_widths):
        for cell in col.cells:
            cell.width = width

    set_table_borders(header_table)

    # Add the logo to the first cell
    cell = header_table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(logo2_path, width=Inches(1.0))

    set_cell_text(header_table.cell(0, 1), '\n\nSCIENCE AND TECHNOLOGY RESEARCH INSTITUTE FOR DEFENCE (STRIDE)', font_size=10, bold=True, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6)

    p1 = doc.add_paragraph()
    run1 = p1.add_run("REVIEW OF REQUEST")
    font_settings_header(run1,font_size=Pt(12), bold=True, underline=False)
    p1.paragraph_format.space_before = Pt(12)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table)
    set_column_width(table.columns[0], 4.25)
    set_column_width(table.columns[1], 3.11)

    set_cell_text(table.cell(0, 0), f'Item: {work_title}', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(0, 1), 'Date:', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 0), f'Client: {client}', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table.cell(1, 1), f'Reference No: {lab_work_no}', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.LEFT)


    p2 = doc.add_paragraph()
    run2 = p2.add_run("")
    font_settings_header(run2, font_size=Pt(12), bold=True, underline=False)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table01 = doc.add_table(rows=9, cols=6)
    table01.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table01)
    set_column_width(table01.columns[0], 0.36)
    set_column_width(table01.columns[1], 2.95)
    set_column_width(table01.columns[2], 0.69)
    set_column_width(table01.columns[3], 0.94)
    set_column_width(table01.columns[4], 0.94)
    set_column_width(table01.columns[5], 1.46)

    set_cell_text(table01.cell(0, 0), 'No', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table01.cell(0, 1), 'Required Test', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table01.cell(0, 2), 'Test Method (√/x)', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table01.cell(0, 3), 'Equipment\n\n (√/x)', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table01.cell(0, 4), 'Laboratory Personnel (√/x)', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table01.cell(0, 5), 'Environmental Condition\n (√/x)', bold=False, line_spacing_pt=12, before_spacing_pt=6,
                  after_spacing_pt=6, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Start populating the table based on type_of_test. 1 row or 2 row.
    row_index = 1  # Starting row index for data

    # Check if any value from {I, II, III, IV} is present in type_of_test
    if any(test in type_of_test for test in {'I', 'II', 'III', 'IV'}):
        table01.add_row()
        set_cell_text(table01.cell(row_index, 0), '1.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table01.cell(row_index, 1), 'Qualitative Analysis', line_spacing_pt=11, before_spacing_pt=3,
                      after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        row_index += 1

    # Check if {IV} is present in type_of_test
    if 'IV' in type_of_test:
        table01.add_row()
        set_cell_text(table01.cell(row_index, 0), '2.', line_spacing_pt=11, before_spacing_pt=3, after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(table01.cell(row_index, 1), 'Quantitative Analysis', line_spacing_pt=11, before_spacing_pt=3,
                      after_spacing_pt=3,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)

    p3 = doc.add_paragraph()
    run3 = p3.add_run("Please tick (√):")
    font_settings_header(run3, bold=False, underline=False)
    p3.paragraph_format.space_before = Pt(2)
    p3.paragraph_format.space_after = Pt(0)
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Define the items to add the checkbox too. We do it this way to easily add checkboxes and iterate it once...
    # instead of re-writing several times to unnecessarily make our codes too long. This is why I comment sparingly too.
    items = [
        "Notify customer & proceed to the job",
        "Reject / Feedback to customer",
        "Subcontracting / Feedback to customer"
    ]

    # Add each item with a checkbox
    for item in items:
        # Create a paragraph for the item
        paragraph = doc.add_paragraph()

        # Add a small square checkbox
        checkbox_run = paragraph.add_run()
        font_settings_header(checkbox_run, bold=False, underline=False, font_size=Pt(20))
        checkbox_run.add_text("☐")  # Unicode for an empty checkbox

        # Add a tab space and the text
        text_run = paragraph.add_run(f" {item}")
        font_settings_header(text_run, bold=False, underline=False)  # Apply your font settings

        # Optional: Adjust paragraph spacing
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    p4 = doc.add_paragraph()
    run4 = p4.add_run("Approved by:")
    font_settings_header(run4, bold=False, underline=False)
    p4.paragraph_format.space_before = Pt(2)
    p4.paragraph_format.space_after = Pt(0)
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table02 = doc.add_table(rows=1, cols=2)
    table02.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_borders(table02)
    set_column_width(table02.columns[0], 5.5)
    set_column_width(table02.columns[1], 1.8)

    set_cell_text(table02.cell(0, 0), 'Laboratory Manager/Deputy Laboratory Manager:', bold=False, line_spacing_pt=12, before_spacing_pt=12,
                  after_spacing_pt=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(table02.cell(0, 1), 'Date:', bold=False, line_spacing_pt=12,
                  before_spacing_pt=12,after_spacing_pt=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    footer = doc.sections[0].footer  # Get the first section of the document
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6))  # Access the header of the section
    footer_table.autofit = False

    # Set explicit column widths
    column_widths = [Inches(2.08), Inches(1.44)]  # Adjust as needed
    for col, width in zip(footer_table.columns, column_widths):
        for cell in col.cells:
            cell.width = width

    set_cell_text(footer_table.cell(0, 0), 'STRIDE/LQP7.1/FORM 1',
                  font_size=8, bold=False, line_spacing_pt=12, before_spacing_pt=6,after_spacing_pt=6,alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(footer_table.cell(0, 1), 'Issue No.: 1\nRev No.: 0\nIssue Date: 15/8/2019',
                  font_size=8, bold=False, line_spacing_pt=12, before_spacing_pt=6, after_spacing_pt=6,alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Save the document
    review_request = os.path.join(doc_out_path, f'{client}_review of request.docx')
    doc.save(review_request)
    messagebox.showinfo("Document Created", f"Word document '{review_request}' created successfully.")

# Function to display the table in a new window
def display_table(next_row):
    # Create a new window for displaying the table
    table_window = tk.Tk()
    table_window.title("Data Table")
    table_window.geometry("1000x600")

    # Create table headers
    headers = ["#", "Report Number", "Date Received", "Client", "Work File","Test Form","Review of Request"]
    for col_num, header in enumerate(headers, start=1):
        tk.Label(table_window, text=header, font=("Arial", 10, "bold"), borderwidth=2, relief="solid", width=15).grid(
            row=0, column=col_num)

    # Loop through the added rows list to display only the new entries using added rows
    for i, row_num in enumerate(added_rows, start=1):  # i is the running number for the entries
        report_number = sheet.cell(row=row_num, column=column_indexes["REPORT NUMBER"]).value
        date_received = sheet.cell(row=row_num, column=column_indexes["DATE RECEIVED"]).value
        item = sheet.cell(row=row_num, column=column_indexes["CLIENT"]).value

        # Insert running number, client, date_received, item, and empty column in the table
        tk.Label(table_window, text=i, borderwidth=2, relief="solid", width=15).grid(row=i, column=1)
        tk.Label(table_window, text=report_number, borderwidth=2, relief="solid", width=15).grid(row=i, column=2)
        tk.Label(table_window, text=date_received, borderwidth=2, relief="solid", width=15).grid(row=i, column=3)
        tk.Label(table_window, text=item, borderwidth=2, relief="solid", width=15).grid(row=i, column=4)

        # Button on 5th, 6th and 7th to generate doc
        button = tk.Button(table_window, text="Work File", command=lambda rn=row_num: create_page1(rn))
        button.grid(row=i, column=5)
        button = tk.Button(table_window, text="Test Form", command=lambda rn=row_num: create_page2(rn))
        button.grid(row=i, column=6)
        button = tk.Button(table_window, text="Review of Request", command=lambda rn=row_num: create_page3(rn))
        button.grid(row=i, column=7)

    table_window.mainloop()


# Create Main Tkinter window for data entry
root = tk.Tk()
root.title("Buku Daftar Makmal")
root.geometry("800x600")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

# Report Number
tk.Label(frame, text="Report Number *").grid(row=0, column=0, padx=5, pady=5, sticky="e")
entry_RN = tk.Entry(frame)
entry_RN.grid(row=0, column=1, padx=5, pady=5)

# Client
tk.Label(frame, text="Client *").grid(row=1, column=0, padx=5, pady=5, sticky="e")
entry_client = tk.Entry(frame)
entry_client.grid(row=1, column=1, padx=5, pady=5)

# Contact Person
tk.Label(frame, text="Contact Person *").grid(row=2, column=0, padx=5, pady=5, sticky="e")
entry_CP = tk.Entry(frame)
entry_CP.grid(row=2, column=1, padx=5, pady=5)

# Work Title
tk.Label(frame, text="Work Title *").grid(row=3, column=0, padx=5, pady=5, sticky="e")
entry_worktitle = tk.Entry(frame)
entry_worktitle.grid(row=3, column=1, padx=5, pady=5)

# Date Received
tk.Label(frame, text="Date Received **").grid(row=4, column=0, padx=5, pady=5, sticky="e")
entry_date_received = tk.Entry(frame)
entry_date_received.grid(row=4, column=1, padx=5, pady=5)

# Received by
tk.Label(frame, text="Received By *").grid(row=5, column=0, padx=5, pady=5, sticky="e")
entry_Rby = tk.Entry(frame)
entry_Rby.grid(row=5, column=1, padx=5, pady=5)

# Quantity
tk.Label(frame, text="Kuantiti *").grid(row=6, column=0, padx=5, pady=5, sticky="e")
entry_kuantiti = tk.Entry(frame)
entry_kuantiti.grid(row=6, column=1, padx=5, pady=5)

# Work Class
tk.Label(frame, text="Work Class *").grid(row=7, column=0, padx=5, pady=5, sticky="e")

workclass_var = tk.StringVar()  # Variable to store selected Work Class
workclass_dropdown = ttk.Combobox(frame, textvariable=workclass_var, state="readonly")
workclass_dropdown["values"] = ["MINDEF", "Berbayar", "Agensi Kerajaan", "STRIDE"]  # Work Class options
workclass_dropdown.grid(row=7, column=1, padx=5, pady=5)

# Binder for event. Only active on so on and forth condition etc2
workclass_dropdown.bind("<<ComboboxSelected>>", on_workclass_subgroup)

tk.Label(frame, text="Sub-Workclass (Mindef @ Berbayar)").grid(row=8, column=0, padx=5, pady=5, sticky="e")

roman_var = tk.StringVar()  # Variable to store selected additional data
roman_dropdown = ttk.Combobox(frame, textvariable=roman_var, state="disabled")  # Initially disabled
roman_dropdown.grid(row=8, column=1, padx=5, pady=5)

tk.Label(frame, text="Mindef Subgroup (D, U, L, MAB)").grid(row=9, column=0, padx=5, pady=5, sticky="e")

additional_var = tk.StringVar()
additional_dropdown = ttk.Combobox(frame, textvariable=additional_var, state="disabled")  # Initially disabled
additional_dropdown.grid(row=9, column=1, padx=5, pady=5)

tk.Label(frame, text="Applicant By").grid(row=10, column=0, padx=5, pady=5, sticky="e")
entry_applicantby = tk.Entry(frame)
entry_applicantby.grid(row=10, column=1, padx=5, pady=5)

#Type of testing checkboxes
microscopic_var = tk.IntVar()
burning_var = tk.IntVar()
solubility_var = tk.IntVar()
quantitative_var = tk.IntVar()
ftir_var = tk.IntVar()

typeoftest_label = tk.Label(frame, text="Type of Testing").grid(row=0, column=2, padx=5, pady=5, sticky="w")
tk.Checkbutton(frame, text="Qualitative Analysis (Microscopic)", variable=microscopic_var).grid(row=1, column=2, padx=5, pady=5, sticky="w")
tk.Checkbutton(frame, text="Qualitative Analysis (Burning)", variable=burning_var).grid(row=2, column=2, padx=5, pady=5, sticky="w")
tk.Checkbutton(frame, text="Qualitative Analysis (Solubility)", variable=solubility_var).grid(row=3, column=2, padx=5, pady=5, sticky="w")
tk.Checkbutton(frame, text="Quantitative Analysis", variable=quantitative_var).grid(row=4, column=2, padx=5, pady=5, sticky="w")
tk.Checkbutton(frame, text="FTIR", variable=ftir_var).grid(row=5, column=2, padx=5, pady=5, sticky="w")


# Sample Marking
tk.Label(frame, text="Sample Marking *").grid(row=6, column=2, padx=5, pady=5, sticky="w")
entry_sample_marking = tk.Entry(frame, width=40)
entry_sample_marking.grid(row=7, column=2, padx=5, pady=5,sticky="w")

instruction_label = tk.Label(frame, text="* Medan wajib diisi, letakkan 'NA' sekiranya tiada data" , fg="blue")
instruction_label.grid(row=14, column=0, columnspan=3, pady=10, sticky=tk.W)
instruction_label = tk.Label(frame, text="** format tarikh ialah DD/MM/YY", fg="blue")
instruction_label.grid(row=15, column=0, columnspan=3, pady=10, sticky=tk.W)
instruction_label = tk.Label(frame, text="Bagi Sample Marking, masukkan item diasingkan dengan ','. Kuantiti sekiranya ada diasingkan dengan ';' {eg: Baju No. 3, Seluar Hitam;2} ", fg="blue")
instruction_label.grid(row=16, column=0, columnspan=3, pady=10, sticky=tk.W)

# Save button
empty_space = tk.Label(frame, text="")
empty_space.grid(row=11, column=0, columnspan=4, pady=10, sticky=tk.W)
save_button = tk.Button(frame, text="Save Data", command=save_data)
save_button.grid(row=12, columnspan=3)

# Status label
label_status = tk.Label(frame, text="")
label_status.grid(row=13, column=0, columnspan=3)

root.mainloop()

